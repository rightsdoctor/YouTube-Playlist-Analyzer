import streamlit as st
import json
import re
import os
import glob
import subprocess
import shutil
import zipfile
import threading
import base64
import pandas as pd
from datetime import datetime
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed

# ============================================================
# 페이지 설정
# ============================================================
st.set_page_config(page_title="YT Playlist Scraper", layout="wide")
st.title("YouTube Playlist Scraper")
st.caption("플레이리스트 URL → 메타데이터 + 자막 → Excel / CSV / 자막 파일")

# ============================================================
# 헬퍼 함수
# ============================================================
INTERNAL_FORMAT = "srt"
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SUBTITLE_DIR = os.path.join(BASE_DIR, "subtitles_temp")
CONVERTED_DIR = os.path.join(BASE_DIR, "subtitles_converted")


def check_ffmpeg_available() -> bool:
    try:
        subprocess.run(["ffmpeg", "-version"],
                       capture_output=True, text=True, timeout=5)
        return True
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def srt_to_plain_text(srt_content: str) -> str:
    lines = srt_content.strip().split('\n')
    text_lines = []
    for line in lines:
        line = line.strip()
        if re.match(r'^\d+$', line):
            continue
        if re.match(r'\d{2}:\d{2}:\d{2}', line):
            continue
        line = re.sub(r'<[^>]+>', '', line)
        if line:
            text_lines.append(line)
    deduplicated = []
    for t in text_lines:
        if not deduplicated or t != deduplicated[-1]:
            deduplicated.append(t)
    return ' '.join(deduplicated)


def format_duration(seconds):
    if not seconds:
        return ''
    seconds = int(seconds)
    h, remainder = divmod(seconds, 3600)
    m, s = divmod(remainder, 60)
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m}:{s:02d}"


def read_subtitle_files(video_id, subtitle_dir, search_exts):
    """
    ★ 수정: video_id 뒤에 반드시 '.'이 오는 패턴만 매칭.
    search_exts: 검색할 확장자 리스트 (예: ["srt", "vtt"])
    """
    if isinstance(search_exts, str):
        search_exts = [search_exts]

    result = {}
    for ext in search_exts:
        # ★ 핵심: "VIDEO_ID.lang.ext" 패턴만 정확히 매칭
        # video_id 뒤에 반드시 '.'이 와야 함 — 다른 영상 ID에 오매칭 방지
        pattern = os.path.join(subtitle_dir, f"{video_id}.*.{ext}")
        sub_files = glob.glob(pattern)

        # 만약 "VIDEO_ID.ext" (언어코드 없는) 형태도 있을 수 있으므로
        exact_file = os.path.join(subtitle_dir, f"{video_id}.{ext}")
        if os.path.exists(exact_file) and exact_file not in sub_files:
            sub_files.append(exact_file)

        for fpath in sub_files:
            fname = os.path.basename(fpath)
            # "VIDEO_ID.ko.srt" → ["VIDEO_ID", "ko", "srt"]
            # "VIDEO_ID.srt" → ["VIDEO_ID", "srt"]
            parts = fname.split(".")
            if len(parts) >= 3:
                # VIDEO_ID.lang.ext → lang은 가운데 부분
                lang = parts[-2]
            else:
                lang = "unknown"
            with open(fpath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read().strip()
                if content:  # ★ 빈 파일 제외
                    result[lang] = content
    return result


def zip_directory(dir_path, ext):
    buf = BytesIO()
    matched = glob.glob(os.path.join(dir_path, f"*.{ext}"))
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for fpath in matched:
            zf.write(fpath, os.path.basename(fpath))
    return buf.getvalue(), len(matched)


def make_download_link(data: bytes, filename: str, label: str) -> str:
    b64 = base64.b64encode(data).decode()
    return (
        f'<a href="data:application/octet-stream;base64,{b64}" '
        f'download="{filename}" '
        f'style="display:inline-block;padding:0.5rem 1rem;'
        f'background-color:#FF4B4B;color:white;text-decoration:none;'
        f'border-radius:0.5rem;font-weight:600;text-align:center;'
        f'width:100%;box-sizing:border-box;">'
        f'{label}</a>'
    )


def count_unique_videos_with_subs(subtitle_dir, exts):
    """자막 디렉토리에서 고유 video_id 수를 정확히 카운트"""
    video_ids_with_subs = set()
    if isinstance(exts, str):
        exts = [exts]
    for ext in exts:
        for fpath in glob.glob(os.path.join(subtitle_dir, f"*.{ext}")):
            fname = os.path.basename(fpath)
            # 첫 번째 '.' 이전이 video_id
            vid = fname.split('.')[0]
            if vid and os.path.getsize(fpath) > 0:
                video_ids_with_subs.add(vid)
    return video_ids_with_subs


# ============================================================
# 사이드바
# ============================================================
with st.sidebar:
    st.header("설정")
    playlist_url = st.text_input(
        "플레이리스트 URL",
        placeholder="https://www.youtube.com/playlist?list=..."
    )

    st.subheader("자막 옵션")
    sub_mode = st.radio(
        "자막 수집 방식",
        ["수동 자막만", "자동 생성 자막만", "수동 우선, 없으면 자동", "수동 + 자동 모두"],
        index=2,
    )
    sub_mode_map = {
        "수동 자막만": "1",
        "자동 생성 자막만": "2",
        "수동 우선, 없으면 자동": "3",
        "수동 + 자동 모두": "4",
    }
    sub_choice = sub_mode_map[sub_mode]
    sub_lang = st.text_input("자막 언어", value="ko",
                             help="예: ko, en, ja 또는 all")
    output_format = st.selectbox("자막 파일 포맷", ["txt", "srt", "vtt", "docx"])
    run_btn = st.button("수집 시작", type="primary", use_container_width=True)

# ============================================================
# 세션 상태 초기화
# ============================================================
if 'collected' not in st.session_state:
    st.session_state.collected = False
    st.session_state.df = None
    st.session_state.errors = []
    st.session_state.csv_data = None
    st.session_state.csv_name = ""
    st.session_state.xlsx_data = None
    st.session_state.xlsx_name = ""
    st.session_state.zip_data = None
    st.session_state.zip_count = 0
    st.session_state.zip_name = ""
    st.session_state.zip_format = ""

# ============================================================
# 메인 실행
# ============================================================
if run_btn and playlist_url:

    st.session_state.collected = False

    has_ffmpeg = check_ffmpeg_available()
    if not has_ffmpeg:
        st.warning("⚠️ ffmpeg 미설치. `packages.txt`에 `ffmpeg`를 추가하세요.")

    for d in [SUBTITLE_DIR, CONVERTED_DIR]:
        if os.path.exists(d):
            shutil.rmtree(d)
        os.makedirs(d, exist_ok=True)

    with st.status("수집 중...", expanded=True) as status:

        # ── 1단계: 영상 ID 수집 ──
        st.write("플레이리스트 분석 중...")
        result = subprocess.run(
            ["yt-dlp", "--flat-playlist", "--dump-json",
             "--no-warnings", "--ignore-errors", playlist_url],
            capture_output=True, text=True, timeout=600,
        )
        flat_entries = []
        for line in result.stdout.strip().split('\n'):
            if line.strip():
                try:
                    flat_entries.append(json.loads(line))
                except json.JSONDecodeError:
                    continue
        video_ids = [e.get('id') or e.get('url', '') for e in flat_entries]
        st.write(f"**{len(video_ids)}개** 영상 감지")

        # ── 2단계: 개별 영상 수집 (병렬) ──
        st.write("개별 영상 메타데이터 + 자막 수집 중... (병렬 처리)")
        progress = st.progress(0)
        full_entries = []
        errors = []
        lock = threading.Lock()
        completed_count = 0

        def process_video(idx, vid):
            url = f"https://www.youtube.com/watch?v={vid}"
            entry = None
            error = None

            try:
                res_meta = subprocess.run(
                    ["yt-dlp", "--skip-download", "--dump-json",
                     "--no-warnings", "--ignore-errors", url],
                    capture_output=True, text=True, timeout=60,
                )
                if res_meta.stdout.strip():
                    entry = json.loads(res_meta.stdout.strip().split('\n')[0])
                    entry['_playlist_position'] = idx
            except Exception as e:
                error = {'position': idx, 'video_id': vid,
                         'error': f"meta: {str(e)}"}
                return entry, error

            # --- 자막 ---
            sub_args = [
                "yt-dlp", "--skip-download",
                "--no-warnings", "--ignore-errors",
                "--write-subs",
                "-o", os.path.join(SUBTITLE_DIR, "%(id)s.%(ext)s"),
            ]

            if has_ffmpeg:
                sub_args += ["--convert-subs", INTERNAL_FORMAT]

            if sub_choice == "1":
                sub_args += ["--no-write-auto-subs"]
            elif sub_choice == "2":
                sub_args.remove("--write-subs")
                sub_args += ["--write-auto-subs"]
            elif sub_choice in ("3", "4"):
                sub_args += ["--write-auto-subs"]

            if sub_lang.lower() == "all":
                sub_args += ["--sub-langs", "all,-live_chat"]
            else:
                langs = [l.strip() for l in sub_lang.split(',')]
                expanded = []
                for l in langs:
                    expanded.append(l)
                    expanded.append(f"{l}-*")
                lang_str = ','.join(expanded) + ',-live_chat'
                sub_args += ["--sub-langs", lang_str]

            sub_args.append(url)

            try:
                subprocess.run(
                    sub_args, capture_output=True, text=True, timeout=120
                )
            except Exception:
                pass

            return entry, error

        total = len(video_ids)

        with ThreadPoolExecutor(max_workers=5) as executor:
            futures = {
                executor.submit(process_video, idx, vid): (idx, vid)
                for idx, vid in enumerate(video_ids, 1)
            }
            for future in as_completed(futures):
                entry, error = future.result()
                with lock:
                    if entry:
                        full_entries.append(entry)
                    if error:
                        errors.append(error)
                    completed_count += 1
                    progress.progress(
                        completed_count / total,
                        text=f"[{completed_count}/{total}] 완료"
                    )

        full_entries.sort(key=lambda x: x.get('_playlist_position', 0))
        progress.progress(1.0, text="수집 완료!")

        # ★ 자막 파일 확장자 결정
        if has_ffmpeg:
            sub_exts_to_search = [INTERNAL_FORMAT]
        else:
            sub_exts_to_search = ["srt", "vtt", "srv1", "srv2", "srv3",
                                  "ttml", "ass", "json3", "lrc"]

        # ★ 고유 video_id 기준으로 자막 보유 영상 수 카운트
        vids_with_subs = count_unique_videos_with_subs(
            SUBTITLE_DIR, sub_exts_to_search
        )

        # 파일 수도 별도 집계 (ZIP용)
        all_sub_files = []
        for ext in sub_exts_to_search:
            all_sub_files.extend(
                glob.glob(os.path.join(SUBTITLE_DIR, f"*.{ext}"))
            )
        # 빈 파일 제거
        all_sub_files = [f for f in all_sub_files if os.path.getsize(f) > 0]

        # 실제 사용할 내부 확장자
        actual_internal_ext = INTERNAL_FORMAT
        if not has_ffmpeg and all_sub_files:
            actual_internal_ext = os.path.splitext(
                all_sub_files[0]
            )[1].lstrip('.')

        st.write(f"자막 파일 **{len(all_sub_files)}개** 수집됨 "
                 f"(영상 **{len(vids_with_subs)}개**)")

        # ── 3단계: txt/docx 변환 ──
        # ★ 핵심 수정: 영상 단위로 변환 (파일 단위가 아닌)
        final_sub_dir = SUBTITLE_DIR
        final_sub_ext = actual_internal_ext
        converted_count = 0

        if output_format in ("txt", "docx") and all_sub_files:
            st.write(f"{output_format.upper()} 변환 중...")
            if output_format == "docx":
                from docx import Document
                from docx.shared import Pt

            # ★ 영상별로 그룹화하여 변환
            # video_id → [파일경로, ...]
            vid_to_files = {}
            for fpath in all_sub_files:
                fname = os.path.basename(fpath)
                vid_from_file = fname.split('.')[0]
                if vid_from_file not in vid_to_files:
                    vid_to_files[vid_from_file] = []
                vid_to_files[vid_from_file].append(fpath)

            for vid_from_file, fpaths in vid_to_files.items():
                # 모든 언어의 자막을 하나의 파일로 합침
                all_plain_parts = []
                for fpath in fpaths:
                    fname = os.path.basename(fpath)
                    parts = fname.split(".")
                    lang = parts[-2] if len(parts) >= 3 else "unknown"
                    with open(fpath, 'r', encoding='utf-8',
                              errors='replace') as f:
                        raw = f.read()
                    plain = srt_to_plain_text(raw)
                    if plain.strip():
                        if len(fpaths) > 1:
                            all_plain_parts.append(f"[{lang}]\n{plain}")
                        else:
                            all_plain_parts.append(plain)

                if not all_plain_parts:
                    continue

                combined_plain = '\n\n'.join(all_plain_parts)
                matched_entry = next(
                    (e for e in full_entries if e.get('id') == vid_from_file),
                    {}
                )
                title = matched_entry.get('title', vid_from_file)
                safe_name = re.sub(
                    r'[^\w가-힣\s]', '', title
                )[:50].strip()

                if output_format == "txt":
                    out_path = os.path.join(
                        CONVERTED_DIR,
                        f"{vid_from_file}_{safe_name}.txt"
                    )
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(f"제목: {title}\n")
                        f.write(f"영상: https://www.youtube.com/watch?v="
                                f"{vid_from_file}\n")
                        f.write(f"{'=' * 60}\n\n")
                        f.write(combined_plain)
                    converted_count += 1

                elif output_format == "docx":
                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.size = Pt(10)
                    style.paragraph_format.line_spacing = 1.5
                    doc.add_heading(title, level=1)
                    meta_p = doc.add_paragraph()
                    meta_p.add_run("영상: ").bold = True
                    meta_p.add_run(
                        f"https://www.youtube.com/watch?v={vid_from_file}"
                    )
                    doc.add_paragraph('─' * 40)
                    sentences = combined_plain.split('. ')
                    buffer = []
                    for s in sentences:
                        buffer.append(s.strip())
                        if len(buffer) >= 4:
                            doc.add_paragraph('. '.join(buffer) + '.')
                            buffer = []
                    if buffer:
                        doc.add_paragraph('. '.join(buffer))
                    out_path = os.path.join(
                        CONVERTED_DIR,
                        f"{vid_from_file}_{safe_name}.docx"
                    )
                    doc.save(out_path)
                    converted_count += 1

            final_sub_dir = CONVERTED_DIR
            final_sub_ext = output_format
            st.write(f"변환 완료: **{converted_count}개**")

        status.update(
            label=f"수집 완료: {len(full_entries)}개 영상",
            state="complete"
        )

    # ── 4단계: DataFrame ──
    rows = []
    for entry in full_entries:
        vid = entry.get('id', '')
        srt_dict = read_subtitle_files(
            vid, SUBTITLE_DIR, sub_exts_to_search
        )
        subtitle_plain = {
            lang: srt_to_plain_text(c) for lang, c in srt_dict.items()
        }
        # ★ 빈 텍스트인 언어는 제거
        subtitle_plain = {
            lang: text for lang, text in subtitle_plain.items()
            if text.strip()
        }

        manual_subs = (
            list(entry.get('subtitles', {}).keys())
            if entry.get('subtitles') else []
        )
        auto_subs = (
            list(entry.get('automatic_captions', {}).keys())
            if entry.get('automatic_captions') else []
        )
        chapters = entry.get('chapters', [])
        chapters_str = ' | '.join(
            [f"{format_duration(ch.get('start_time', 0))} "
             f"{ch.get('title', '')}" for ch in chapters]
        ) if chapters else ''
        thumbnails = entry.get('thumbnails', [])
        best_thumb = thumbnails[-1].get('url', '') if thumbnails else ''

        row = {
            '#': entry.get('_playlist_position', ''),
            'video_url': f"https://www.youtube.com/watch?v={vid}",
            'video_id': vid,
            'title': entry.get('title', ''),
            'description': entry.get('description', ''),
            'channel': entry.get('channel', ''),
            'channel_id': entry.get('channel_id', ''),
            'channel_url': entry.get('channel_url', ''),
            'uploader': entry.get('uploader', ''),
            'channel_follower_count': entry.get('channel_follower_count', ''),
            'upload_date': entry.get('upload_date', ''),
            'view_count': entry.get('view_count', ''),
            'like_count': entry.get('like_count', ''),
            'comment_count': entry.get('comment_count', ''),
            'duration_seconds': entry.get('duration', ''),
            'duration_readable': format_duration(entry.get('duration')),
            'categories': ', '.join(entry.get('categories') or []),
            'tags': ', '.join(entry.get('tags') or []),
            'language': entry.get('language', ''),
            'age_limit': entry.get('age_limit', ''),
            'live_status': entry.get('live_status', ''),
            'availability': entry.get('availability', ''),
            'thumbnail_url': best_thumb,
            'chapters': chapters_str,
            'manual_subtitle_langs':
                ', '.join(manual_subs[:30]) if manual_subs else '',
            'auto_subtitle_langs':
                ', '.join(auto_subs[:15]) if auto_subs else '',
            'subtitle_collected_langs': ', '.join(subtitle_plain.keys()),
        }
        for lang, text in subtitle_plain.items():
            row[f'subtitle_text_{lang}'] = text
        rows.append(row)

    if rows:
        df = pd.DataFrame(rows)
    else:
        df = pd.DataFrame(columns=[
            '#', 'video_url', 'video_id', 'title', 'description', 'channel',
            'channel_id', 'channel_url', 'uploader', 'channel_follower_count',
            'upload_date', 'view_count', 'like_count', 'comment_count',
            'duration_seconds', 'duration_readable', 'categories', 'tags',
            'language', 'age_limit', 'live_status', 'availability',
            'thumbnail_url', 'chapters', 'manual_subtitle_langs',
            'auto_subtitle_langs', 'subtitle_collected_langs',
        ])

    # ── 다운로드 데이터를 session_state에 저장 ──
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    st.session_state.csv_data = df.to_csv(
        index=False, encoding='utf-8-sig'
    ).encode('utf-8-sig')
    st.session_state.csv_name = f"playlist_{timestamp}.csv"

    xlsx_buf = BytesIO()
    with pd.ExcelWriter(xlsx_buf, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Videos')
        sub_cols = [c for c in df.columns if c.startswith('subtitle_text_')]
        if sub_cols:
            df[['#', 'video_id', 'title'] + sub_cols].to_excel(
                writer, index=False, sheet_name='Subtitles')
    st.session_state.xlsx_data = xlsx_buf.getvalue()
    st.session_state.xlsx_name = f"playlist_{timestamp}.xlsx"

    zip_data, zip_count = zip_directory(final_sub_dir, final_sub_ext)
    st.session_state.zip_data = zip_data if zip_count > 0 else None
    st.session_state.zip_count = zip_count
    st.session_state.zip_name = f"subtitles_{output_format}_{timestamp}.zip"
    st.session_state.zip_format = output_format

    st.session_state.df = df
    st.session_state.errors = errors
    st.session_state.collected = True

# ============================================================
# 결과 표시 & 다운로드
# ============================================================
if (st.session_state.collected
        and st.session_state.df is not None
        and not st.session_state.df.empty):
    df = st.session_state.df
    errors = st.session_state.errors

    # ★ 핵심 수정: metric 계산을 실제 자막 텍스트 존재 여부로 정확히 판단
    sub_text_cols = [c for c in df.columns if c.startswith('subtitle_text_')]
    if sub_text_cols:
        # 자막 텍스트 컬럼 중 하나라도 비어있지 않은 행
        has_sub_mask = df[sub_text_cols].apply(
            lambda row: any(
                str(v).strip() != '' and str(v).strip() != 'nan'
                for v in row
            ),
            axis=1
        )
        sub_count = has_sub_mask.sum()
    elif 'subtitle_collected_langs' in df.columns:
        sub_count = (
            df['subtitle_collected_langs'].astype(str)
            .apply(lambda x: x.strip() != '' and x.strip() != 'nan')
        ).sum()
    else:
        sub_count = 0

    c1, c2, c3 = st.columns(3)
    c1.metric("총 영상", f"{len(df)}개")
    c2.metric("자막 수집", f"{sub_count}개")
    c3.metric("실패", f"{len(errors)}개")

    display_cols = [
        '#', 'title', 'channel', 'duration_readable',
        'view_count', 'like_count', 'subtitle_collected_langs'
    ]
    display_cols = [c for c in display_cols if c in df.columns]

    st.dataframe(df[display_cols], use_container_width=True, height=400)

    st.subheader("다운로드")
    d1, d2, d3 = st.columns(3)

    with d1:
        st.markdown(
            make_download_link(
                st.session_state.csv_data,
                st.session_state.csv_name,
                "CSV"
            ),
            unsafe_allow_html=True,
        )

    with d2:
        st.markdown(
            make_download_link(
                st.session_state.xlsx_data,
                st.session_state.xlsx_name,
                "XLSX"
            ),
            unsafe_allow_html=True,
        )

    with d3:
        if st.session_state.zip_data:
            st.markdown(
                make_download_link(
                    st.session_state.zip_data,
                    st.session_state.zip_name,
                    f"자막 ZIP ({st.session_state.zip_format}, "
                    f"{st.session_state.zip_count}개)"
                ),
                unsafe_allow_html=True,
            )

    if errors:
        with st.expander(f"실패 로그 ({len(errors)}건)"):
            st.dataframe(pd.DataFrame(errors))

elif run_btn and not playlist_url:
    st.warning("플레이리스트 URL을 입력하세요.")
